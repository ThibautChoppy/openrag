"""DOCX ``DocumentParser`` implementation.

Conversion to Markdown via the ``markitdown`` library; falls back to
plain-text extraction via ``python-docx`` if MarkItDown fails.

MarkItDown emits a generic ``![](data:image/png;base64...)`` placeholder
(literal, truncated) for every embedded image, with no per-image
identifier. Actual image bytes are pulled from the DOCX zip
(``word/media/``) and matched to placeholders **positionally**, in
document order. Each placeholder is rewritten to a unique synthetic
``![](docx-image-N)`` ref, and the matching :class:`ImageBlock` stores
that ref in ``metadata['markdown_ref']`` for downstream caption
substitution.

Captioning is not done here — see :class:`ImageBlock` for the
parser→caption contract.

Output:
- A single ``TextBlock`` containing the rewritten Markdown.
- One ``ImageBlock`` per embedded zip image; ``caption=None``.

Failures fall back gracefully: missing libraries or malformed zips
degrade to leaving content untouched rather than raising. Images that
can't be decoded by PIL (e.g. EMF, WMF) are skipped — the matching
placeholder in the markdown is left in place for downstream cleanup.
"""

from __future__ import annotations

import asyncio
import logging
import re
import zipfile
from io import BytesIO

from ...models.document import Document, DocumentType, ImageBlock, ProcessedDocument, TextBlock
from ..image_preprocessor import ensure_png_compatible_mode, pil_to_png_bytes
from .document_parser import DocumentParser
from .registry import parser_registry

logger = logging.getLogger(__name__)


# Match MarkItDown's image refs in the rendered markdown. The current
# version emits a truncated placeholder (``![](data:image/png;base64...)``)
# but older / future versions may emit a full data URI or non-empty alt
# text. The pattern below matches both shapes — same regex the legacy
# loader used (``components/indexer/loaders/docx.py``).
_MARKITDOWN_IMAGE_PLACEHOLDER = re.compile(r"!\[.*?\]\(data:image/[^)]*\)")


def _image_ref(index: int) -> str:
    """Synthetic markdown image ref used as a placeholder for embedded DOCX images."""
    return f"![](docx-image-{index})"


@parser_registry.register("docx")
class DocxParser(DocumentParser):
    """Parse DOCX into a Markdown TextBlock + one ImageBlock per embedded image."""

    def supported_types(self) -> list[str]:
        return [DocumentType.DOCX.value]

    async def parse(self, document: Document) -> ProcessedDocument:
        if not document.raw_bytes:
            return ProcessedDocument(
                document_id=document.id,
                metadata=dict(document.metadata),
            )

        async with document.as_temporary_file() as path:
            markdown, embedded = await asyncio.to_thread(self._extract, str(path))

        markdown, images = self._rewrite_placeholders_and_build_blocks(markdown, embedded)
        markdown = markdown.strip()
        text_blocks = [TextBlock(text=markdown, page_number=1)] if markdown else []
        return ProcessedDocument(
            document_id=document.id,
            text_blocks=text_blocks,
            images=images,
            metadata=dict(document.metadata),
            page_count=1 if markdown else 0,
        )

    # ----- helpers -----

    @classmethod
    def _extract(cls, path: str) -> tuple[str, list[bytes | None]]:
        """Run MarkItDown + zip-image extraction in one thread hop."""
        return cls._convert_to_markdown(path), cls._extract_embedded_images(path)

    @staticmethod
    def _convert_to_markdown(path: str) -> str:
        try:
            from markitdown import MarkItDown
        except ImportError:
            logger.warning("markitdown not available; falling back to python-docx text extraction")
            return DocxParser._fallback_extract_text(path)
        try:
            return MarkItDown().convert(path).text_content
        except Exception as exc:
            logger.warning("MarkItDown DOCX conversion failed (%s); falling back to plain text", exc)
            return DocxParser._fallback_extract_text(path)

    @staticmethod
    def _fallback_extract_text(path: str) -> str:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("python-docx not available; cannot extract DOCX text")
            return ""
        try:
            doc = DocxDocument(path)
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            logger.warning("python-docx fallback failed: %s", exc)
            return ""

    @staticmethod
    def _extract_embedded_images(path: str) -> list[bytes | None]:
        """Return PNG bytes for each embedded image in document order.

        ``None`` entries preserve positional alignment with markdown
        placeholders for unsupported formats (EMF, WMF, …).
        """
        try:
            from PIL import Image
        except ImportError:
            logger.warning("PIL not available; skipping DOCX image extraction")
            return []
        try:
            with zipfile.ZipFile(path, "r") as zf:
                media = [n for n in zf.namelist() if n.startswith("word/media/")]
                if not media:
                    return []
                ordered: dict[int, bytes | None] = {}
                for name in media:
                    raw = zf.read(name)
                    try:
                        order_num = int(name.split("media/image")[1].split(".")[0])
                    except (IndexError, ValueError):
                        continue
                    try:
                        with Image.open(BytesIO(raw)) as im:
                            im = ensure_png_compatible_mode(im)
                            ordered[order_num] = pil_to_png_bytes(im)
                    except Exception as exc:
                        logger.warning("Skipping unsupported DOCX media %s: %s", name, exc)
                        ordered[order_num] = None
                if not ordered:
                    return []
                max_order = max(ordered)
                return [ordered.get(i + 1) for i in range(max_order)]
        except zipfile.BadZipFile:
            logger.warning("DOCX is not a valid zip archive; skipping image extraction")
            return []
        except Exception as exc:
            logger.warning("DOCX image extraction failed: %s", exc)
            return []

    @staticmethod
    def _rewrite_placeholders_and_build_blocks(
        markdown: str, embedded: list[bytes | None]
    ) -> tuple[str, list[ImageBlock]]:
        """Replace each MarkItDown placeholder with a unique synthetic ref
        and emit one ``ImageBlock`` per successfully-decoded zip image.

        Positional matching: the i-th placeholder in the markdown maps to
        the i-th entry in ``embedded``. ``None`` entries (unsupported
        formats) collapse the placeholder to an empty string.
        """
        if not markdown or not embedded:
            return markdown, []

        images: list[ImageBlock] = []
        idx = 0
        consumed = 0  # count of `embedded` entries used so far

        def replacer(_match: re.Match[str]) -> str:
            nonlocal idx, consumed
            if consumed >= len(embedded):
                return ""  # more placeholders than zip images: drop extras
            payload = embedded[consumed]
            consumed += 1
            if payload is None:
                return ""  # zip image was unsupported; drop placeholder
            ref = _image_ref(idx)
            idx += 1
            images.append(
                ImageBlock(
                    image_bytes=payload,
                    page_number=1,
                    mime_type="image/png",
                    metadata={"markdown_ref": ref},
                )
            )
            return ref

        new_markdown = _MARKITDOWN_IMAGE_PLACEHOLDER.sub(replacer, markdown)
        return new_markdown, images
